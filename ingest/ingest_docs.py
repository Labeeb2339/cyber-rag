#!/usr/bin/env python3
"""
CyberRAG document ingester — add YOUR OWN documents to the knowledge base.

This is the feature that makes CyberRAG usable by a real SOC/CERT: drop your
sensitive PDFs, Word docs, and text files into a folder and ingest them. Nothing
leaves the machine — extraction, chunking, embedding (nomic-embed-text via
Ollama) and storage (ChromaDB on disk) are all 100% local.

Supports: .pdf (PyMuPDF), .docx (python-docx), .md/.txt (plain), .html (stripped).

INCREMENTAL by design — it ADDS to the existing `cyber_kb` collection without
wiping the authoritative corpus. Re-ingesting the same file replaces its chunks
(stable IDs), so you can re-run safely after a document is updated.

Usage:
  python ingest/ingest_docs.py <path-to-file-or-folder> [--source my-org] [--recursive]
  python ingest/ingest_docs.py ./my_reports --source internal-cti --recursive

After ingesting, rebuild BM25 so the new docs are searchable by exact term:
  python rag/hybrid.py build
"""
import os
import re
import sys
import glob
import argparse
import hashlib

import chromadb
import ollama
from langchain_text_splitters import RecursiveCharacterTextSplitter

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHROMA_DIR = os.path.join(ROOT, "data", "chroma")
COLLECTION = "cyber_kb"
EMBED_MODEL = "nomic-embed-text"

ATTACK_RE = re.compile(r"\bT\d{4}(?:\.\d{3})?\b")
CVE_RE = re.compile(r"\bCVE-\d{4}-\d{4,7}\b", re.I)
SUPPORTED = {".pdf", ".docx", ".md", ".txt", ".markdown", ".html", ".htm"}
_TAG = re.compile(r"<[^>]+>")


def extract_text(path):
    """Return plain text for a supported document, or None if unreadable."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            text = "\n\n".join(page.get_text("text") for page in doc)
            doc.close()
            return text
        if ext == ".docx":
            import docx
            d = docx.Document(path)
            paras = [p.text for p in d.paragraphs]
            # also pull table cell text — advisories love tables
            for t in d.tables:
                for row in t.rows:
                    paras.append(" | ".join(c.text for c in row.cells))
            return "\n".join(paras)
        if ext in (".html", ".htm"):
            with open(path, encoding="utf-8", errors="ignore") as f:
                return _TAG.sub(" ", f.read())
        # .md / .txt / .markdown
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  [skip] {path}: {e}")
        return None


def embed(texts):
    return [ollama.embeddings(model=EMBED_MODEL, prompt=t)["embedding"] for t in texts]


def gather(target, recursive):
    if os.path.isfile(target):
        return [target]
    pat = "**/*" if recursive else "*"
    files = []
    for p in glob.glob(os.path.join(target, pat), recursive=recursive):
        if os.path.isfile(p) and os.path.splitext(p)[1].lower() in SUPPORTED:
            files.append(p)
    return sorted(files)


def main():
    ap = argparse.ArgumentParser(description="Ingest your own docs into CyberRAG (local-only).")
    ap.add_argument("target", help="file or folder to ingest")
    ap.add_argument("--source", default="user-docs",
                    help="source label stored in metadata (default: user-docs)")
    ap.add_argument("--recursive", action="store_true", help="recurse into subfolders")
    args = ap.parse_args()

    files = gather(args.target, args.recursive)
    if not files:
        print(f"No supported documents found at {args.target}")
        print(f"Supported: {', '.join(sorted(SUPPORTED))}")
        sys.exit(1)
    print(f"[ingest-docs] {len(files)} document(s), source='{args.source}'")

    client = chromadb.PersistentClient(path=CHROMA_DIR)
    col = client.get_or_create_collection(COLLECTION, metadata={"hnsw:space": "cosine"})

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1800, chunk_overlap=200,
        separators=["\n## ", "\n### ", "\n\n", "\n", ". ", " "],
    )

    total_chunks = 0
    for n, path in enumerate(files, 1):
        text = extract_text(path)
        if not text or not text.strip():
            print(f"  [{n}/{len(files)}] empty/unreadable: {os.path.basename(path)}")
            continue

        attack_ids = sorted(set(ATTACK_RE.findall(text)))
        cve_ids = sorted(set(c.upper() for c in CVE_RE.findall(text)))
        meta_base = {
            "source": args.source,
            "doc": os.path.basename(path),
            "path": path,
            "attack_ids": ",".join(attack_ids[:10]),
            "cve_ids": ",".join(cve_ids[:10]),
        }

        chunks = splitter.split_text(text)
        if not chunks:
            continue

        # stable IDs keyed on path -> re-ingesting a file overwrites its old chunks
        ids = [hashlib.md5(f"{path}:{j}".encode()).hexdigest() for j in range(len(chunks))]
        metas = [{**meta_base, "chunk": j} for j in range(len(chunks))]

        # upsert in batches to bound memory
        B = 64
        for s in range(0, len(chunks), B):
            sl = slice(s, s + B)
            col.upsert(ids=ids[sl], embeddings=embed(chunks[sl]),
                       documents=chunks[sl], metadatas=metas[sl])
        total_chunks += len(chunks)
        tags = []
        if attack_ids:
            tags.append(f"{len(attack_ids)} ATT&CK")
        if cve_ids:
            tags.append(f"{len(cve_ids)} CVE")
        tag = f"  ({', '.join(tags)})" if tags else ""
        print(f"  [{n}/{len(files)}] {os.path.basename(path)} -> {len(chunks)} chunks{tag}")

    print(f"[ingest-docs] DONE — {total_chunks} chunks added/updated")
    print(f"[ingest-docs] collection now holds {col.count()} chunks total")
    print("[ingest-docs] NEXT: rebuild BM25 so new docs are term-searchable ->  python rag/hybrid.py build")


if __name__ == "__main__":
    main()
